import streamlit as st
import pandas as pd
from docx import Document
import random
import io
import traceback
from zipfile import ZipFile
import re
import unicodedata

st.set_page_config(page_title="Générateur de QCM", layout="centered")
st.title("Générateur de QCM personnalisés")


# — Fonctions utilitaires —

def remplacer_placeholder_dans_paragraphe(paragraph, replacements):
    """
    Remplace dans tout le texte du paragraphe (paragraph.text)
    les placeholders définis dans `replacements`, même s'ils sont
    fragmentés en plusieurs runs au sein d'un même paragraphe.
    """
    if not paragraph.text:
        return

    # Texte brut du paragraphe
    texte_complet = paragraph.text
    for key, value in replacements.items():
        # Trois formes : espace normal, espace insécable, ou pas d'espace
        ni = key.replace(" ", "\u00a0")
        ns = key.replace(" ", "")
        texte_complet = texte_complet.replace(key, value)
        texte_complet = texte_complet.replace(ni, value)
        texte_complet = texte_complet.replace(ns, value)

    # Supprime tous les runs existants
    for run in paragraph.runs:
        run.text = ""
    # Insère un nouveau run avec le texte modifié
    paragraph.add_run(texte_complet)


def detecter_questions(doc):
    """
    Renvoie une liste de questions détectées dans `doc` (python-docx.Document).
    Chaque question est un dict :
      {
        "index": int,            # index du paragraphe où commence la question
        "texte": str,            # libellé “X.Y - Texte ?” ou “n - Texte ?”
        "numero": str,           # numéro “X.Y” nettoyé ou séquentiel “n”
        "reponses": [            # liste de dicts pour chaque réponse
            {
              "index": int,      # index du paragraphe où se trouve cette réponse
              "lettre": str,     # “A”|"B"|"C"|"D"
              "texte": str,      # texte de la réponse
              "correct": bool,   # True si {{checkbox}} est présent
              "original_text": str
            }, …
        ],
        "correct_idx": int,      # index dans “reponses” de la bonne réponse
        "original_text": str     # texte complet du paragraphe de la question
      }
    Ne conserve que les questions ayant ≥ 2 réponses et au moins une réponse correcte.
    """
    questions = []
    current_question = None

    pattern_num = re.compile(
        r'^\s*'                    # début de ligne, espaces optionnels
        r'(\d+(?:\s*\.\s*\d+)*)'   # capture “1.2” ou “1 . 2” etc.
        r'\s*'                     # espaces optionnels
        r'[-–—.]?'                 # un tiret (court, long) ou un point (optionnel)
        r'\s*'                     # espaces optionnels
        r'(.+?)'                   # texte de la question
        r'\s*\?$'                  # “?” en fin de ligne
    )
    pattern_non_num = re.compile(r'^\s*[-–—]\s*(.+?)\s*\?$')

    for i, para in enumerate(doc.paragraphs):
        raw = para.text.strip()
        if not raw:
            continue
        texte = raw.replace("\u00a0", " ").replace("–", "-").replace("—", "-")

        # 1) Question numérotée
        m_num = pattern_num.match(texte)
        if m_num:
            num_brut = m_num.group(1)
            num = re.sub(r'\s*\.\s*', '.', num_brut).strip()
            txt = m_num.group(2).strip()
            libelle = f"{num} - {txt}?"
            current_question = {
                "index": i,
                "texte": libelle,
                "numero": num,
                "reponses": [],
                "correct_idx": None,
                "original_text": raw
            }
            questions.append(current_question)
            continue

        # 2) Question non-numérotée
        m_non = pattern_non_num.match(texte)
        if m_non:
            seq = str(len(questions) + 1)
            txt = m_non.group(1).strip()
            libelle = f"{seq} - {txt}?"
            current_question = {
                "index": i,
                "texte": libelle,
                "numero": seq,
                "reponses": [],
                "correct_idx": None,
                "original_text": raw
            }
            questions.append(current_question)
            continue

        # 3) Réponse A–D pour la question en cours
        if current_question:
            m_ans = re.match(r'^([A-D])\s*[-\s–—.]+\s*(.*?)\s*(\{\{checkbox\}\})?$', texte, re.IGNORECASE)
            if m_ans:
                lettre = m_ans.group(1).upper()
                rep_txt = m_ans.group(2).strip()
                is_corr = bool(m_ans.group(3))
                current_question["reponses"].append({
                    "index": i,
                    "lettre": lettre,
                    "texte": rep_txt,
                    "correct": is_corr,
                    "original_text": raw
                })
                if is_corr:
                    current_question["correct_idx"] = len(current_question["reponses"]) - 1

    valid = [q for q in questions if q["correct_idx"] is not None and len(q["reponses"]) >= 2]
    for q in questions:
        if q not in valid:
            st.warning(f"Ignorée : {q['texte']} (moins de 2 réponses ou pas de {{checkbox}})")
    return valid


def parse_correct_answers(f):
    """
    Lit un fichier Excel comportant au moins deux colonnes :
      - 'Numéro de la question' (ex. '1.1', '2.3', '3', etc.)
      - 'Réponse correcte'      (ex. 'A', 'B', 'C' ou 'D')
    Retourne dict {question_num: lettre_correcte}.
    """
    if f is None:
        return {}
    try:
        # Forcer la colonne "Numéro de la question" en texte
        df = pd.read_excel(f, dtype={'Numéro de la question': str})
        df = df.dropna(subset=['Numéro de la question', 'Réponse correcte'])
        df['Numéro de la question'] = df['Numéro de la question'].astype(str).str.strip()
        df['Réponse correcte'] = df['Réponse correcte'].astype(str).str.strip().str.upper()

        # Vérifier que les réponses correctes ne contiennent que A, B, C ou D
        if not df['Réponse correcte'].isin(['A', 'B', 'C', 'D']).all():
            mauvaises_valeurs = df.loc[~df['Réponse correcte'].isin(['A', 'B', 'C', 'D']), 'Réponse correcte'].unique()
            st.error(f"Valeurs invalides dans 'Réponse correcte' : {mauvaises_valeurs}")
            return {}

        return dict(zip(df['Numéro de la question'], df['Réponse correcte']))
    except Exception as e:
        st.error(f"Erreur lecture corrections : {e}")
        return {}


def calculer_resultat_final(score, total_q):
    """
    Renvoie l'étiquette selon le pourcentage (score/total_q) * 100 :
      ≥75% → 'Acquis'
      50–75% → 'En cours d’acquisition'
      <50%  → 'Non acquis'
    """
    if total_q <= 0:
        return "Non acquis"
    pct = (score / total_q) * 100
    if pct >= 75:
        return "Acquis"
    elif pct >= 50:
        return "En cours d’acquisition"
    else:
        return "Non acquis"


def slugify(value):
    """
    Normalise une chaîne pour la rendre sûre dans un nom de fichier :
    suppression des accents, conversion en ASCII, et remplacement des espaces par des underscores.
    """
    value = unicodedata.normalize('NFKD', str(value))
    value = value.encode('ascii', 'ignore').decode('ascii')
    value = re.sub(r'[^\w\s-]', '', value).strip().lower()
    return re.sub(r'[-\s]+', '_', value)


def generer_document(row, template_bytes):
    """
    Génère un Document .docx pour un apprenant donné (ligne `row` de l'Excel).
    Retourne (Document, score_total, résultat_final, total_questions).
    Utilise st.session_state['questions'] et st.session_state['correct_answers'].
    """
    try:
        doc = Document(io.BytesIO(template_bytes))

        # --- 1) Remplacement des placeholders “apprenant” ---
        date_eval = row['Date Évaluation']
        try:
            date_eval_dt = pd.to_datetime(date_eval, dayfirst=True)
            date_eval = date_eval_dt.strftime("%d/%m/%Y")
        except Exception:
            date_eval = str(date_eval)
            st.warning(f"Format inattendu pour la date d’évaluation de {row['Prénom']} {row['Nom']} : « {date_eval} »")

        repl_apprenant = {
            '{{prenom}}': str(row['Prénom']),
            '{{nom}}': str(row['Nom']),
            '{{email}}': str(row['Email']),
            '{{ref_session}}': str(row['Référence Session']),
            '{{date_evaluation}}': date_eval
        }
        # Remplacer dans tous les paragraphes et tables
        for p in doc.paragraphs:
            remplacer_placeholder_dans_paragraphe(p, repl_apprenant)
        for tbl in doc.tables:
            for r in tbl.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        remplacer_placeholder_dans_paragraphe(p, repl_apprenant)

        # --- 2) Préparer le comptage par module ---
        questions = st.session_state['questions']
        questions_par_module = {}
        correct_par_module = {}
        total_questions = len(questions)

        for q in questions:
            if '.' in q['numero']:
                module_key = q['numero'].split('.')[0]
            else:
                module_key = q['numero']
            questions_par_module[module_key] = questions_par_module.get(module_key, 0) + 1
            correct_par_module.setdefault(module_key, 0)

        corr = st.session_state.get('correct_answers', {})
        score_total = 0

        # --- 3) Traiter chaque question (placements et score) ---
        for q in questions:
            if '.' in q['numero']:
                module_key = q['numero'].split('.')[0]
            else:
                module_key = q['numero']

            reps = q['reponses'].copy()
            q_num = q['numero']

            if st.session_state['figees'].get(q['index'], False):
                chosen_idx = st.session_state['reponses_correctes'].get(q['index'], q['correct_idx'])
                bonne_rep = reps.pop(chosen_idx)
                reps.insert(0, bonne_rep)
            else:
                if q['correct_idx'] is not None:
                    # Extraire la bonne réponse
                    bonne_rep = reps.pop(q['correct_idx'])
                    # Mélanger uniquement les autres réponses
                    autres_reps = reps[:]
                    random.shuffle(autres_reps)
                    reps = [bonne_rep] + autres_reps

            for r in reps:
                idx_para = r['index']
                if idx_para < len(doc.paragraphs):
                    box = "☑" if reps.index(r) == 0 else "☐"
                    doc.paragraphs[idx_para].clear()
                    doc.paragraphs[idx_para].add_run(f"{r['lettre']} - {r['texte']} {box}")

            # Calcul du score si une correction est connue
            if q_num in corr:
                generated_answer = reps[0]['lettre'].upper()
                expected_answer = corr[q_num].upper()
                if generated_answer == expected_answer:
                    correct_par_module[module_key] += 1
                    score_total += 1

        # --- 4) Préparer les remplacements “scores par module” ---
        sr = {}
        for module_key, tot in questions_par_module.items():
            score_mod = correct_par_module.get(module_key, 0)
            sr[f'{{{{result_mod{module_key}}}}}'] = str(score_mod)
            sr[f'{{{{total_mod{module_key}}}}}'] = str(tot)

        sr['{{result_mod_total}}'] = str(score_total)
        sr['{{total_questions}}'] = str(total_questions)
        sr['{{result_evaluation}}'] = calculer_resultat_final(score_total, total_questions)

        # Remplacer dans tout le document (paragraphes, tables, en-têtes, pieds de page)
        for p in doc.paragraphs:
            remplacer_placeholder_dans_paragraphe(p, sr)
        for tbl in doc.tables:
            for r in tbl.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        remplacer_placeholder_dans_paragraphe(p, sr)
        for section in doc.sections:
            for header_para in section.header.paragraphs:
                remplacer_placeholder_dans_paragraphe(header_para, sr)
            for footer_para in section.footer.paragraphs:
                remplacer_placeholder_dans_paragraphe(footer_para, sr)

        return doc, score_total, sr['{{result_evaluation}}'], total_questions

    except Exception as e:
        st.error(f"Erreur génération doc : {e}")
        st.error(traceback.format_exc())
        return None, 0, "Erreur", 0


# — Interface Streamlit —

with st.expander("Étape 1 : Importation des fichiers", expanded=True):
    excel_file = st.file_uploader(
        "Excel (Prénom, Nom, Email, Référence Session, Date Évaluation)",
        type="xlsx"
    )
    word_file = st.file_uploader("Modèle Word (.docx)", type="docx")
    corr_file = st.file_uploader("Réponses correctes (Excel .xlsx)", type="xlsx")


# Initialiser le session_state si nécessaire
if 'questions' not in st.session_state:
    st.session_state['questions'] = []
if 'figees' not in st.session_state:
    st.session_state['figees'] = {}
if 'reponses_correctes' not in st.session_state:
    st.session_state['reponses_correctes'] = {}
if 'current_template' not in st.session_state:
    st.session_state['current_template'] = None
if 'correct_answers' not in st.session_state:
    st.session_state['correct_answers'] = {}


# 1) Charger le modèle Word et détecter les questions
if word_file and st.session_state['current_template'] != word_file.name:
    try:
        data = word_file.getvalue()
        doc_sample = Document(io.BytesIO(data))
        qs = detecter_questions(doc_sample)
        st.session_state['questions'] = qs
        st.session_state['current_template'] = word_file.name
        st.session_state['figees'] = {}
        st.session_state['reponses_correctes'] = {}

        if qs:
            st.success(f"✅ {len(qs)} questions détectées")
            with st.expander("🔍 Questions détectées", expanded=True):
                for idx, q in enumerate(qs, 1):
                    st.write(f"**{idx}. {q['texte']}**")
                    for j, r in enumerate(q['reponses']):
                        mark = "✅" if j == q['correct_idx'] else "☐"
                        st.write(f"{mark} {r['lettre']} - {r['texte']}")
        else:
            st.warning("⚠️ Aucune question détectée. Vérifiez le format du Word.")
    except Exception as e:
        st.error(f"Erreur chargement Word : {e}")
        st.error(traceback.format_exc())


# 2) Charger le fichier des corrections
if corr_file:
    ca = parse_correct_answers(corr_file)
    st.session_state['correct_answers'] = ca
    st.success(f"✅ {len(ca)} corrections chargées")


# 3) Configuration des questions “figées”
if st.session_state['questions']:
    st.markdown("### Configuration des questions")
    for q in st.session_state['questions']:
        q_id = q['index']
        q_num = q['numero']
        key_base = f"{q_id}_{st.session_state['current_template']}"
        fig = st.checkbox(f"Figer question {q_num}", key=f"figer_{key_base}")
        if fig:
            options = [f"{r['lettre']} - {r['texte']}" for r in q['reponses']]
            default_idx = q['correct_idx'] if q['correct_idx'] is not None else 0
            choix = st.selectbox(
                f"Choix pour {q_num}",
                options=options,
                index=default_idx,
                key=f"bonne_{key_base}"
            )
            st.session_state['figees'][q_id] = True
            st.session_state['reponses_correctes'][q_id] = options.index(choix)
        else:
            st.session_state['figees'].pop(q_id, None)
            st.session_state['reponses_correctes'].pop(q_id, None)


# 4) Génération des QCM
if excel_file and word_file and st.session_state['questions']:
    if st.button("Générer les QCM"):
        try:
            df = pd.read_excel(excel_file)
            required_cols = ['Prénom', 'Nom', 'Email', 'Référence Session', 'Date Évaluation']
            missing = [c for c in required_cols if c not in df.columns]
            if missing:
                st.error(f"Colonnes manquantes dans l'Excel : {missing}")
                st.stop()

            buf = io.BytesIO()
            recap = []

            with ZipFile(buf, 'w') as zf:
                progress = st.progress(0)
                total = len(df)
                for i, row in df.iterrows():
                    doc_out, sc, res, tot_q = generer_document(row, word_file.getvalue())
                    if doc_out:
                        recap.append({
                            "Prénom": row["Prénom"],
                            "Nom": row["Nom"],
                            "Réf": row["Référence Session"],
                            "Score": sc,
                            "Total Questions": tot_q,
                            "Pourcentage": f"{(sc/tot_q)*100:.1f}%" if tot_q > 0 else "0%",
                            "Résultat": res
                        })
                        bytes_io = io.BytesIO()
                        doc_out.save(bytes_io)
                        fn = f"QCM_{slugify(row['Prénom'])}_{slugify(row['Nom'])}_{slugify(row['Référence Session'])}.docx"
                        zf.writestr(fn, bytes_io.getvalue())
                    progress.progress((i + 1) / total)

                if recap:
                    df_r = pd.DataFrame(recap)
                    excel_buf = io.BytesIO()
                    with pd.ExcelWriter(excel_buf, engine='openpyxl') as writer:
                        df_r.to_excel(writer, index=False, sheet_name="Récapitulatif")
                    excel_buf.seek(0)
                    zf.writestr("Recapitulatif_QCM.xlsx", excel_buf.getvalue())

            buf.seek(0)
            st.success("✅ Génération terminée")
            st.download_button(
                "⬇️ Télécharger l’archive ZIP",
                data=buf,
                file_name="QCM_Personnalises.zip",
                mime="application/zip"
            )

            st.subheader("Récapitulatif des résultats")
            st.dataframe(pd.DataFrame(recap))

        except Exception as e:
            st.error(f"ERREUR critique : {e}")
            st.error(traceback.format_exc())


# 5) Légende des résultats
st.markdown("### Légende résultats")
st.info("""
- **Acquis** : ≥ 75%  
- **En cours d’acquisition** : 50–75%  
- **Non acquis** : < 50%
""")
